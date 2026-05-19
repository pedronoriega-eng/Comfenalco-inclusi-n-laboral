# -*- coding: utf-8 -*-
# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  MOTOR DOCUMENTAL — Generador de Informe APT                              ║
# ║  Módulo refactorizado para integración web (Streamlit)                     ║
# ║  Programa Talento Sin Barreras — Comfenalco Santander                     ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

import os
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DatosProyecto:
    """
    Contenedor dinámico de TODOS los datos del informe.
    Acepta un diccionario de configuración para poblar cada campo.
    """

    # Valores por defecto (ejemplo Mas por Menos)
    DEFAULTS = {
        "empresa": [
            "Supermercados Mas por Menos SAS",
            "Comercio al por menor en establecimientos no especializados, "
            "con surtido compuesto principalmente por alimentos, bebidas y tabaco",
            "900.119.072",
            "Bucaramanga",
            "Por confirmar con la empresa",
            "Por confirmar con la empresa",
            "Mayo de 2026",
        ],
        "cargo_encabezado": [
            "Vendedor / Surtidor – Oficios Varios",
            "Operativo",
            "Punto de Venta – Sala de Ventas",
            "Jefe de Zona",
        ],
        "cargo_objetivo": (
            "Apoyar de forma integral la operación del punto de venta en el área de surtido, "
            "dando un servicio con calidad y calidez de acuerdo con los protocolos establecidos "
            "por la compañía."
        ),
        "cargo_requisitos": [
            "Básica secundaria",
            "No se requieren certificaciones adicionales",
            "Servicio al cliente, surtido, manejo de inventarios",
            "Experiencia en logística, despachos, servicio al cliente y ventas",
            "Entrenamiento en el puesto de trabajo con acompañamiento del Jefe de Zona",
            "Tiempo completo, disponibilidad para turnos rotativos",
            "Exámenes médicos ocupacionales de ingreso y periódicos "
            "según el profesiograma de la empresa",
        ],
        "cargo_tareas": [
            "1. Surtir y resurtir los productos agotados en el piso de venta.",
            "2. Asegurarse de que las exhibiciones estén bien abastecidas.",
            "3. Confirmar que las etiquetas coincidan con los precios.",
            "4. Examinar la fecha de vencimiento de los productos.",
            "5. Frentear la mercancía ubicada en los lineales.",
            "6. Desarmar las cajas de cartón y almacenarlas.",
            "7. Almacenar cajas vacías correctamente.",
            "8. Cumplir con la misión de venta y promover productos.",
            "9. Gestionar la devolución de implementos.",
            "10. Llenar formato de re-empaque.",
            "11. Reportar productos faltantes.",
            "12. Ubicar el material P.O.P.",
            "13. Seguir las normas de servicio al cliente.",
        ],
        "condiciones_org": [
            "Lunes a sábado, jornada diurna conforme a la Ley 2466 de 2025.",
            "Se contemplan turnos de apertura y cierre del punto de venta.",
            "Turnos rotativos según programación del Jefe de Zona.",
            "Rotación entre secciones del punto de venta según necesidad.",
            "Eventualmente se pueden presentar horas extras conforme a la ley.",
        ],
        "recursos": [
            "Computador de punto de venta, lectores de código de barras, básculas digitales.",
            "Estanterías (lineales), góndolas, exhibidores, mesas de trabajo.",
            "No aplica maquinaria pesada para este cargo.",
            "Gato hidráulico (estibador manual), escaleras tipo tijera, carros transportadores.",
            "Bolsas de empaque, etiquetas de precios, material P.O.P., formatos impresos.",
            "Uniformes de dotación institucional.",
            "Guantes de protección para manipulación de productos, calzado antideslizante.",
        ],
        "accesibilidad": [
            "Pendiente de análisis visual por A-Vision.",
            "",
            "Pendiente de análisis visual por A-Vision.",
            "",
            "Pendiente de análisis visual por A-Vision.",
            "No se cuenta con registro fotográfico.",
            "Pendiente de análisis visual por A-Vision.",
            "Pendiente de análisis visual por A-Vision.",
            "Pendiente de análisis visual por A-Vision.",
        ],
        "condiciones_amb": [
            "Pendiente de análisis.",
            "Pendiente de análisis.",
            "Pendiente de análisis.",
            "Pendiente de análisis.",
            "Pendiente de análisis.",
        ],
        "condiciones_seguridad": [
            "Pendiente de análisis.",
            "No aplica",
            "Pendiente de análisis.",
            "Matriz de identificación de peligros (GTC 45)",
        ],
        "competencias": [
            # INTELECTUALES (8)
            "Medio", "Medio", "Medio", "Bajo", "Medio", "Bajo", "Bajo", "Alto",
            # PERSONALES (2)
            "Alto", "Medio",
            # INTERPERSONALES (6)
            "Alto", "Alto", "Bajo", "Medio", "Alto", "Medio",
            # ORGANIZACIONALES (4)
            "Medio", "Alto", "Medio", "Medio",
            # TECNOLÓGICAS (2)
            "Bajo", "Bajo",
            # EMPRESARIALES (5)
            "Bajo", "Bajo", "Medio", "Alto", "Bajo",
            # FÍSICAS (8)
            "Bajo", "Alto", "Bajo", "Alto", "Medio", "Alto", "Bajo", "Bajo",
        ],
        "discapacidades_sugeridas": [
            "Auditiva",
            "Intelectual Leve",
        ],
        "ajustes_razonables": [
            {
                "tipo": "Auditiva",
                "fisicas": "Pendiente de análisis por A-Legal.",
                "comunicacion": "Pendiente de análisis por A-Legal.",
                "actitudinales": "Pendiente de análisis por A-Legal.",
            },
            {
                "tipo": "Intelectual Leve",
                "fisicas": "Pendiente de análisis por A-Legal.",
                "comunicacion": "Pendiente de análisis por A-Legal.",
                "actitudinales": "Pendiente de análisis por A-Legal.",
            },
        ],
        "conclusiones": "Pendiente de generación por el sistema multiagente.",
    }

    def __init__(self, config: dict = None):
        """
        Inicializa DatosProyecto con los valores proporcionados o defaults.

        Args:
            config: Diccionario con las claves de datos del informe.
                    Las claves no proporcionadas usarán los valores por defecto.
        """
        cfg = {**self.DEFAULTS, **(config or {})}

        # Rutas de archivos
        self.PLANTILLA_APT = cfg.get("plantilla_path", "")
        self.ARCHIVO_SALIDA = cfg.get("salida_path", "")
        self.FOTOS = cfg.get("fotos", {})
        self.FOTO_ANCHO_PULGADAS = cfg.get("foto_ancho", 3.5)

        # Datos de tablas
        self.EMPRESA = cfg["empresa"]
        self.CARGO_ENCABEZADO = cfg["cargo_encabezado"]
        self.CARGO_OBJETIVO = cfg["cargo_objetivo"]
        self.CARGO_REQUISITOS = cfg["cargo_requisitos"]
        self.CARGO_TAREAS = cfg["cargo_tareas"]
        self.CONDICIONES_ORG = cfg["condiciones_org"]
        self.RECURSOS = cfg["recursos"]
        self.ACCESIBILIDAD = cfg["accesibilidad"]
        self.CONDICIONES_AMB = cfg["condiciones_amb"]
        self.CONDICIONES_SEGURIDAD = cfg["condiciones_seguridad"]
        self.COMPETENCIAS = cfg["competencias"]
        self.DISCAPACIDADES_SUGERIDAS = cfg["discapacidades_sugeridas"]
        self.AJUSTES_RAZONABLES = cfg["ajustes_razonables"]
        self.CONCLUSIONES = cfg["conclusiones"]


class MotorDocumental:
    """
    Agente A-Doc: Motor de inyección de datos en plantilla Word.
    Preserva formato original (fuentes, bordes, alineaciones).
    Adaptado para integración web con callbacks de progreso.
    """

    PLACEHOLDER = "[[REF_IA_X]]"

    def __init__(self, datos: DatosProyecto, progress_callback=None):
        """
        Args:
            datos: Instancia de DatosProyecto con todos los datos.
            progress_callback: Función callback(fase: str, progreso: float)
                              fase = nombre de la etapa actual
                              progreso = 0.0 a 1.0
        """
        self.d = datos
        self.doc = None
        self.tables = None
        self.stats = {"reemplazos": 0, "fotos": 0, "errores": []}
        self._progress = progress_callback or (lambda fase, prog: None)
        self._logs = []

    # ── Utilidades de reemplazo ──────────────────────────────────────

    @staticmethod
    def _reemplazar_en_celda(cell, old_text, new_text):
        """Reemplaza placeholder en celda preservando formato del primer run."""
        for paragraph in cell.paragraphs:
            if old_text in paragraph.text:
                if paragraph.runs:
                    first_run = paragraph.runs[0]
                    for run in paragraph.runs:
                        run.text = ""
                    first_run.text = new_text
                    return True
        return False

    @staticmethod
    def _limpiar_runs_celda(cell):
        """Limpia todos los runs de una celda."""
        first_run = None
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if first_run is None:
                    first_run = run
                run.text = ""
        return first_run

    @staticmethod
    def _insertar_imagen_en_celda(cell, ruta_imagen, ancho_pulgadas=3.5):
        """Inserta imagen en celda, centrando y limpiando texto previo."""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(ruta_imagen, width=Inches(ancho_pulgadas))

    def _reemplazar_en_parrafos_cuerpo(self, old_text, new_text):
        """Reemplaza placeholder en párrafos fuera de tablas."""
        for paragraph in self.doc.paragraphs:
            if old_text in paragraph.text:
                if paragraph.runs:
                    for run in paragraph.runs:
                        run.text = ""
                    paragraph.runs[0].text = new_text
                    self.stats["reemplazos"] += 1
                    return True
        return False

    # ── Procesamiento por tabla ──────────────────────────────────────

    def _procesar_tabla_simple(self, tabla_idx, valores, col_valor=1, fila_inicio=0):
        tabla = self.tables[tabla_idx]
        for i, valor in enumerate(valores):
            fila = fila_inicio + i
            try:
                cell = tabla.rows[fila].cells[col_valor]
                if self._reemplazar_en_celda(cell, self.PLACEHOLDER, valor):
                    self.stats["reemplazos"] += 1
            except IndexError:
                self.stats["errores"].append(f"T{tabla_idx} R{fila} C{col_valor}: fuera de rango")

    def _procesar_tabla_1_celda(self, tabla_idx, valores, fila_inicio, col=0):
        tabla = self.tables[tabla_idx]
        for i, valor in enumerate(valores):
            fila = fila_inicio + i
            try:
                cell = tabla.rows[fila].cells[col]
                if self._reemplazar_en_celda(cell, self.PLACEHOLDER, valor):
                    self.stats["reemplazos"] += 1
            except IndexError:
                self.stats["errores"].append(f"T{tabla_idx} R{fila} C{col}: fuera de rango")

    # ── Flujo principal ──────────────────────────────────────────────

    def generar(self):
        """Ejecuta el flujo completo de generación del informe."""
        self._log("═" * 60)
        self._log("  GENERADOR DE INFORME TÉCNICO DE INCLUSIÓN LABORAL")
        self._log("  Programa Talento Sin Barreras — Comfenalco Santander")
        self._log("═" * 60)

        # [1/8] Cargar plantilla
        self._progress("Abriendo plantilla Word...", 0.05)
        self._log("[1/8] Abriendo plantilla Word...")
        self.doc = Document(self.d.PLANTILLA_APT)
        self.tables = self.doc.tables
        self._log(f"      Tablas encontradas: {len(self.tables)}")

        # [2/8] Datos de la empresa
        self._progress("Datos de la empresa...", 0.15)
        self._log("[2/8] Datos de la empresa (Tabla 0)...")
        self._procesar_tabla_simple(0, self.d.EMPRESA, col_valor=1, fila_inicio=0)

        # [3/8] Perfil de cargo
        self._progress("Perfil de cargo y tareas...", 0.25)
        self._log("[3/8] Perfil de cargo y tareas (Tabla 1)...")
        self._procesar_tabla_simple(1, self.d.CARGO_ENCABEZADO, col_valor=1, fila_inicio=0)
        self._procesar_tabla_1_celda(1, [self.d.CARGO_OBJETIVO], fila_inicio=5, col=0)
        self._procesar_tabla_simple(1, self.d.CARGO_REQUISITOS, col_valor=1, fila_inicio=7)
        self._procesar_tabla_1_celda(1, self.d.CARGO_TAREAS, fila_inicio=15, col=0)

        # [4/8] Condiciones organizacionales y recursos
        self._progress("Condiciones organizacionales...", 0.35)
        self._log("[4/8] Condiciones organizacionales (Tabla 2)...")
        self._procesar_tabla_simple(2, self.d.CONDICIONES_ORG, col_valor=1, fila_inicio=0)
        self._log("[4/8] Recursos materiales (Tabla 3)...")
        self._procesar_tabla_simple(3, self.d.RECURSOS, col_valor=1, fila_inicio=0)

        # [5/8] Accesibilidad y fotos
        self._progress("Accesibilidad y barreras...", 0.50)
        self._log("[5/8] Accesibilidad, barreras y fotos (Tabla 5)...")
        self._procesar_accesibilidad()

        # [6/8] Condiciones ambientales y competencias
        self._progress("Condiciones ambientales...", 0.60)
        self._log("[6/8] Condiciones ambientales (Tabla 6)...")
        self._procesar_tabla_simple(6, self.d.CONDICIONES_AMB, col_valor=1, fila_inicio=1)
        self._procesar_condiciones_seguridad()

        self._progress("Competencias...", 0.70)
        self._log("[6/8] Competencias (Tabla 7)...")
        self._procesar_competencias()

        # [7/8] Discapacidades y ajustes
        self._progress("Discapacidades y ajustes razonables...", 0.80)
        self._log("[7/8] Discapacidades sugeridas y ajustes (Tablas 8-9)...")
        self._procesar_discapacidades_sugeridas()
        self._procesar_ajustes_razonables()

        # [8/8] Conclusiones
        self._progress("Conclusiones...", 0.90)
        self._log("[8/8] Conclusiones...")
        self._reemplazar_en_parrafos_cuerpo(self.PLACEHOLDER, self.d.CONCLUSIONES)

        # Guardar
        self._progress("Guardando informe final...", 0.95)
        self._log(f"\n      Guardando: {self.d.ARCHIVO_SALIDA}")
        self.doc.save(self.d.ARCHIVO_SALIDA)

        # Verificación
        self._verificar()
        self._progress("¡Informe generado exitosamente!", 1.0)

        return self.stats

    def generar_en_memoria(self):
        """
        Genera el informe y lo devuelve como bytes en memoria.
        Útil para Streamlit st.download_button().
        """
        self._log("═" * 60)
        self._log("  GENERADOR DE INFORME TÉCNICO DE INCLUSIÓN LABORAL")
        self._log("═" * 60)

        # Cargar plantilla
        self._progress("Abriendo plantilla Word...", 0.05)
        self._log("[1/8] Abriendo plantilla Word...")
        self.doc = Document(self.d.PLANTILLA_APT)
        self.tables = self.doc.tables
        self._log(f"      Tablas encontradas: {len(self.tables)}")

        # Procesar todas las tablas
        self._progress("Datos de la empresa...", 0.15)
        self._procesar_tabla_simple(0, self.d.EMPRESA, col_valor=1, fila_inicio=0)

        self._progress("Perfil de cargo...", 0.25)
        self._procesar_tabla_simple(1, self.d.CARGO_ENCABEZADO, col_valor=1, fila_inicio=0)
        self._procesar_tabla_1_celda(1, [self.d.CARGO_OBJETIVO], fila_inicio=5, col=0)
        self._procesar_tabla_simple(1, self.d.CARGO_REQUISITOS, col_valor=1, fila_inicio=7)
        self._procesar_tabla_1_celda(1, self.d.CARGO_TAREAS, fila_inicio=15, col=0)

        self._progress("Condiciones organizacionales...", 0.35)
        self._procesar_tabla_simple(2, self.d.CONDICIONES_ORG, col_valor=1, fila_inicio=0)
        self._procesar_tabla_simple(3, self.d.RECURSOS, col_valor=1, fila_inicio=0)

        self._progress("Accesibilidad y barreras...", 0.50)
        self._procesar_accesibilidad()

        self._progress("Condiciones ambientales...", 0.60)
        self._procesar_tabla_simple(6, self.d.CONDICIONES_AMB, col_valor=1, fila_inicio=1)
        self._procesar_condiciones_seguridad()

        self._progress("Competencias...", 0.70)
        self._procesar_competencias()

        self._progress("Discapacidades y ajustes...", 0.80)
        self._procesar_discapacidades_sugeridas()
        self._procesar_ajustes_razonables()

        self._progress("Conclusiones...", 0.90)
        self._reemplazar_en_parrafos_cuerpo(self.PLACEHOLDER, self.d.CONCLUSIONES)

        # Guardar en memoria
        self._progress("Generando archivo final...", 0.95)
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)

        self._verificar()
        self._progress("¡Informe generado exitosamente!", 1.0)

        return buffer.getvalue(), self.stats

    # ── Procesamiento específico ─────────────────────────────────────

    def _procesar_accesibilidad(self):
        """Tabla 5: Inyecta textos de accesibilidad y fotos."""
        tabla = self.tables[5]

        for i, texto in enumerate(self.d.ACCESIBILIDAD):
            row_idx = i + 1
            try:
                cell = tabla.rows[row_idx].cells[1]
                first_run = self._limpiar_runs_celda(cell)
                if first_run:
                    first_run.text = texto
                    self.stats["reemplazos"] += 1
            except IndexError:
                self.stats["errores"].append(f"T5 R{row_idx}: fuera de rango")

        for fila_idx, ruta_foto in self.d.FOTOS.items():
            if os.path.exists(ruta_foto):
                try:
                    self._insertar_imagen_en_celda(
                        tabla.rows[fila_idx].cells[1],
                        ruta_foto,
                        self.d.FOTO_ANCHO_PULGADAS,
                    )
                    self.stats["fotos"] += 1
                    self._log(f"      ✓ Foto insertada fila {fila_idx}: {os.path.basename(ruta_foto)}")
                except Exception as e:
                    self.stats["errores"].append(f"Error foto fila {fila_idx}: {e}")
            else:
                self.stats["errores"].append(f"Foto no encontrada: {ruta_foto}")

    def _procesar_condiciones_seguridad(self):
        tabla = self.tables[6]
        fila = tabla.rows[6]
        for i, valor in enumerate(self.d.CONDICIONES_SEGURIDAD):
            if self._reemplazar_en_celda(fila.cells[i], self.PLACEHOLDER, valor):
                self.stats["reemplazos"] += 1

    def _procesar_competencias(self):
        tabla = self.tables[7]
        comp_idx = 0
        for row in tabla.rows:
            last_cell = row.cells[-1]
            if self.PLACEHOLDER in last_cell.text and comp_idx < len(self.d.COMPETENCIAS):
                if self._reemplazar_en_celda(last_cell, self.PLACEHOLDER, self.d.COMPETENCIAS[comp_idx]):
                    self.stats["reemplazos"] += 1
                    comp_idx += 1
        self._log(f"      Competencias inyectadas: {comp_idx}/{len(self.d.COMPETENCIAS)}")

    def _procesar_discapacidades_sugeridas(self):
        tabla = self.tables[8]
        for i, disc in enumerate(self.d.DISCAPACIDADES_SUGERIDAS):
            if self._reemplazar_en_celda(tabla.rows[i].cells[0], self.PLACEHOLDER, disc):
                self.stats["reemplazos"] += 1

    def _procesar_ajustes_razonables(self):
        tabla = self.tables[9]
        fila_base = 1
        for ajuste in self.d.AJUSTES_RAZONABLES:
            try:
                if self._reemplazar_en_celda(tabla.rows[fila_base].cells[1], self.PLACEHOLDER, ajuste["tipo"]):
                    self.stats["reemplazos"] += 1
                if self._reemplazar_en_celda(tabla.rows[fila_base + 1].cells[2], self.PLACEHOLDER, ajuste["fisicas"]):
                    self.stats["reemplazos"] += 1
                if self._reemplazar_en_celda(tabla.rows[fila_base + 2].cells[2], self.PLACEHOLDER, ajuste["comunicacion"]):
                    self.stats["reemplazos"] += 1
                if self._reemplazar_en_celda(tabla.rows[fila_base + 3].cells[2], self.PLACEHOLDER, ajuste["actitudinales"]):
                    self.stats["reemplazos"] += 1
            except IndexError:
                self.stats["errores"].append(f"Ajustes T9 fila_base={fila_base}: fuera de rango")
            fila_base += 4

    # ── Verificación ─────────────────────────────────────────────────

    def _verificar(self):
        restantes = 0
        for tabla in self.doc.tables:
            for row in tabla.rows:
                for cell in row.cells:
                    if self.PLACEHOLDER in cell.text:
                        restantes += 1
        for p in self.doc.paragraphs:
            if self.PLACEHOLDER in p.text:
                restantes += 1

        self._log(f"\n      ─── VERIFICACIÓN ───")
        self._log(f"      Reemplazos realizados: {self.stats['reemplazos']}")
        self._log(f"      Fotos insertadas:      {self.stats['fotos']}")
        self._log(f"      Placeholders restantes: {restantes}")

        if self.stats["errores"]:
            self._log(f"      Errores ({len(self.stats['errores'])}):")
            for err in self.stats["errores"]:
                self._log(f"        ✗ {err}")

        if restantes > 0:
            self._log(f"\n      ⚠ ADVERTENCIA: Quedan {restantes} placeholders sin reemplazar.")
        else:
            self._log(f"      ✓ Todos los placeholders fueron reemplazados.")

        self.stats["placeholders_restantes"] = restantes

    # ── Logging ──────────────────────────────────────────────────────

    def _log(self, msg):
        self._logs.append(msg)

    def get_logs(self):
        """Devuelve todos los logs de la generación."""
        return "\n".join(self._logs)


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  API SIMPLIFICADA                                                          ║
# ╚══════════════════════════════════════════════════════════════════════════════╝


def generar_informe(config: dict, progress_callback=None):
    """
    API de alto nivel para generar el informe.

    Args:
        config: Diccionario con todos los datos del informe.
                Debe incluir 'plantilla_path' y 'salida_path'.
        progress_callback: Función callback(fase, progreso) opcional.

    Returns:
        tuple: (bytes del archivo .docx, dict de estadísticas, str de logs)
    """
    datos = DatosProyecto(config)
    motor = MotorDocumental(datos, progress_callback)
    docx_bytes, stats = motor.generar_en_memoria()
    return docx_bytes, stats, motor.get_logs()
