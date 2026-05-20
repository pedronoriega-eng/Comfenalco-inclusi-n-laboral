# -*- coding: utf-8 -*-
# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  SISTEMA MULTIAGENTE DE IA — Análisis de Inclusión Laboral                 ║
# ║  Agentes: A-Vision (Accesibilidad) + A-Legal (Jurídico)                   ║
# ║  Programa Talento Sin Barreras — Comfenalco Santander                     ║
# ╚══════════════════════════════════════════════════════════════════════════════╝
#
# Este módulo integra la API de Google Gemini para ejecutar los agentes
# de análisis visual y jurídico del sistema multiagente.
#
# Requiere: GEMINI_API_KEY configurada como variable de entorno o
#           en Streamlit secrets (st.secrets["GEMINI_API_KEY"])

import json
import os
import base64
import re
import io
from pathlib import Path
from PIL import Image

try:
    import google.generativeai as genai
    GEMINI_DISPONIBLE = True
except ImportError:
    GEMINI_DISPONIBLE = False


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  CONFIGURACIÓN                                                             ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

MODELO_GEMINI = "gemini-2.5-flash"

# Marco legal referenciado por los agentes
MARCO_LEGAL = """
MARCO LEGAL COLOMBIANO APLICABLE:

1. Ley 2466 de 2025 (Reforma Laboral):
   - Art. 20: Estabilidad laboral reforzada para personas con discapacidad.
   - Art. 21: Ajustes razonables obligatorios con el mínimo de modificaciones.
   - Jornada máxima: 42 horas semanales con reducción gradual.

2. Ley 1618 de 2013 (Derechos de PcD):
   - Art. 2: Definición de ajustes razonables y diseño universal.
   - Art. 8: Sensibilización y toma de conciencia.
   - Art. 14: Acceso al entorno físico (rampas, señalización, NTC 4143/6047).
   - Art. 15: Acceso al transporte.
   - Art. 21: Acceso a la información y comunicaciones.

3. Ley 361 de 1997:
   - Art. 26: Protección laboral reforzada, no discriminación.

4. Normas Técnicas:
   - NTC 4143: Rampas fijas (pendiente máx. 12%, ancho mín. 90 cm).
   - NTC 6047: Espacios accesibles (baños, giro mín. 1.50 m).
"""


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  AGENTE A-VISION: Análisis de Accesibilidad                               ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

PROMPT_VISION = """
Eres el Agente A-Vision: Analista de Riesgos y Accesibilidad Arquitectónica.
Tu especialidad es identificar barreras físicas, de tránsito y de acceso para
personas con discapacidad en entornos laborales colombianos.

{marco_legal}

INSTRUCCIONES:
1. Analiza la imagen proporcionada con extremo detalle.
2. Identifica TODAS las barreras arquitectónicas visibles.
3. Describe riesgos de movilidad y accesibilidad.
4. Referencia las normas técnicas colombianas aplicables.
5. Sé preciso y técnico. NO inventes lo que no se ve en la imagen.

TIPO DE ESPACIO: {tipo_espacio}

Responde en español con un párrafo técnico detallado (mínimo 100 palabras)
describiendo los hallazgos de accesibilidad para este espacio.
Incluye referencias legales específicas (artículos de ley).
NO uses formato markdown, solo texto plano en un párrafo continuo.
"""

PROMPT_VISION_BATCH = """
Eres el Agente A-Vision: Experto en Accesibilidad e Inclusión Laboral.
Analiza detenidamente TODAS las imágenes adjuntas, las cuales pertenecen a un mismo entorno laboral.
Para facilitar la referencia, las fotos están en el mismo orden que fueron cargadas, empezando por el índice 0.

{marco_legal}

Tu tarea consta de dos partes:
1. MAPEO FOTOGRÁFICO: Identifica la mejor foto (por su índice, 0 a N) que represente cada una de estas áreas clave: "entrada", "baños", y "cafetería". Si no hay foto para un área, asigna -1.
2. ANÁLISIS INTEGRAL: Genera un análisis técnico detallado consolidando lo visto en TODAS las fotos para cada categoría, referenciando las leyes colombianas (Ley 1618 de 2013, NTC 4143, NTC 6047). Cada análisis debe ser un párrafo continuo de al menos 80 palabras.
Si un área no tiene foto, redacta un análisis técnico preventivo indicando que no hay registro fotográfico pero se sugieren mejoras generales.

Responde estrictamente en formato JSON con la siguiente estructura (sin markdown, sin ```json):
{{
    "mapeo": {{
        "entrada": 0,
        "banos": 2,
        "cafeteria": -1
    }},
    "analisis": {{
        "acceso_principal": "Análisis consolidado de la entrada...",
        "banos": "Análisis consolidado de baños...",
        "cafeteria": "Análisis consolidado de cafetería...",
        "rampas": "Análisis de rampas, pasillos y escaleras...",
        "informacion_comunicacion": "Análisis de señalización y sistemas de comunicación...",
        "actitudinales_sociales": "Análisis de barreras actitudinales y sociales evidentes o preventivas...",
        "condiciones_biologico": "Exposición a riesgos biológicos...",
        "condiciones_psicosocial": "Factores de riesgo psicosocial...",
        "condiciones_biomecanico": "Riesgos biomecánicos (posturas, movimientos, cargas)...",
        "condiciones_fisico": "Condiciones físicas (iluminación, ruido, ventilación)...",
        "condiciones_quimico": "Exposición a sustancias químicas (si no aplica, justificar)...",
        "condiciones_seguridad": "Condiciones de seguridad (riesgo locativo, emergencias)..."
    }}
}}
"""


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  AGENTE A-LEGAL: Diagnóstico Jurídico de Inclusión                        ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

PROMPT_LEGAL = """
Eres el Agente A-Legal: Jurista Especialista en Inclusión Laboral y Talento Humano.
Tu expertise incluye la Ley 2466 de 2025, Ley 1618 de 2013 y Ley 361 de 1997.

{marco_legal}

DATOS DEL CARGO:
- Nombre: {cargo_nombre}
- Tipo: {cargo_tipo}
- Objetivo: {cargo_objetivo}
- Tareas principales:
{tareas}

ANÁLISIS DE ACCESIBILIDAD (del Agente A-Vision):
{analisis_vision}

INSTRUCCIONES:
1. Analiza las funciones del cargo y determina qué tipos de discapacidad son
   COMPATIBLES con las tareas descritas.
2. Para cada tipo de discapacidad compatible, define ajustes razonables que
   prioricen la infraestructura ACTUAL de la empresa (mínimo de modificaciones).
3. Los ajustes deben clasificarse en: Físicas, Comunicación, Actitudinales.
4. Genera conclusiones técnico-jurídicas del informe.

Responde ÚNICAMENTE con el JSON (sin markdown, sin ```json):

{{
    "discapacidades_compatibles": [
        "Tipo 1 (ej: Auditiva)",
        "Tipo 2 (ej: Intelectual Leve)"
    ],
    "ajustes_razonables": [
        {{
            "tipo": "Nombre de la discapacidad",
            "fisicas": "Párrafo detallado de ajustes físicos/arquitectónicos con referencias legales (mín. 80 palabras)...",
            "comunicacion": "Párrafo detallado de ajustes de comunicación con referencias legales (mín. 80 palabras)...",
            "actitudinales": "Párrafo detallado de ajustes actitudinales con referencias legales (mín. 80 palabras)..."
        }}
    ],
    "competencias": [
        "Nivel – Justificación para Atención (INTELECTUAL 1)",
        "Nivel – Justificación para Memoria (INTELECTUAL 2)",
        "Nivel – Justificación para Concentración (INTELECTUAL 3)",
        "Nivel – Justificación para Cálculo (INTELECTUAL 4)",
        "Nivel – Justificación para Solución problemas (INTELECTUAL 5)",
        "Nivel – Justificación para Toma decisiones (INTELECTUAL 6)",
        "Nivel – Justificación para Creatividad (INTELECTUAL 7)",
        "Nivel – Justificación para Orientación espacial (INTELECTUAL 8)",
        "Nivel – Justificación para Ética (PERSONAL 1)",
        "Nivel – Justificación para Adaptación (PERSONAL 2)",
        "Nivel – Justificación para Comunicación (INTERPERSONAL 1)",
        "Nivel – Justificación para Trabajo equipo (INTERPERSONAL 2)",
        "Nivel – Justificación para Liderazgo (INTERPERSONAL 3)",
        "Nivel – Justificación para Manejo conflictos (INTERPERSONAL 4)",
        "Nivel – Justificación para Flexibilidad (INTERPERSONAL 5)",
        "Nivel – Justificación para Proactividad (INTERPERSONAL 6)",
        "Nivel – Justificación para Gestión información (ORGANIZACIONAL 1)",
        "Nivel – Justificación para Servicio cliente (ORGANIZACIONAL 2)",
        "Nivel – Justificación para Gestión recursos (ORGANIZACIONAL 3)",
        "Nivel – Justificación para Resp. ambiental (ORGANIZACIONAL 4)",
        "Nivel – Justificación para Uso tecnología (TECNOLÓGICA 1)",
        "Nivel – Justificación para Innovación (TECNOLÓGICA 2)",
        "Nivel – Justificación para Identificar oportunidades (EMPRESARIAL 1)",
        "Nivel – Justificación para Elaborar planes (EMPRESARIAL 2)",
        "Nivel – Justificación para Referenciación (EMPRESARIAL 3)",
        "Nivel – Justificación para Mercadeo (EMPRESARIAL 4)",
        "Nivel – Justificación para Plan de negocios (EMPRESARIAL 5)",
        "Nivel – Justificación para Posición sedente (FÍSICA 1)",
        "Nivel – Justificación para Posición bípeda (FÍSICA 2)",
        "Nivel – Justificación para Alternancia (FÍSICA 3)",
        "Nivel – Justificación para Mov. repetitivos (FÍSICA 4)",
        "Nivel – Justificación para Desplazamientos (FÍSICA 5)",
        "Nivel – Justificación para Levantamiento cargas (FÍSICA 6)",
        "Nivel – Justificación para Trabajo alturas (FÍSICA 7)",
        "Nivel – Justificación para Espacios confinados (FÍSICA 8)"
    ],
    "conclusiones": "Párrafo completo de conclusiones técnico-jurídicas del informe (mínimo 200 palabras). Debe incluir: compatibilidad del cargo, barreras identificadas, ajustes propuestos, recomendaciones, y referencia al Programa Talento Sin Barreras de Comfenalco Santander."
}}

REGLAS:
- Los niveles de competencia DEBEN ser: Bajo, Medio o Alto, seguidos de un guión y justificación.
- Los ajustes razonables DEBEN priorizar infraestructura actual (mínimo de modificaciones).
- Las conclusiones DEBEN referenciar artículos de ley específicos.
- PROHIBIDO inventar funciones o leyes. Solo usar la información proporcionada.
- Genera exactamente 35 competencias en el orden indicado.
"""


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  CLASE PRINCIPAL: Orquestador Multiagente                                 ║
# ╚══════════════════════════════════════════════════════════════════════════════╝


class SistemaMultiagente:
    """
    Orquesta los agentes A-Vision y A-Legal para generar
    el análisis completo de inclusión laboral.
    """

    def __init__(self, api_key: str):
        """
        Args:
            api_key: Clave de API de Google Gemini.
        """
        if not GEMINI_DISPONIBLE:
            raise ImportError(
                "El paquete 'google-generativeai' no está instalado. "
                "Ejecute: pip install google-generativeai"
            )

        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel(MODELO_GEMINI)
        self._logs = []

    def _log(self, msg):
        self._logs.append(msg)

    def get_logs(self):
        return "\n".join(self._logs)

    # ── A-Vision ─────────────────────────────────────────────────────

    def ejecutar_vision(self, rutas_fotos: list) -> dict:
        """
        Ejecuta el Agente A-Vision sobre las fotos proporcionadas
        procesándolas en un solo lote para evitar límites de API (Rate Limits).
        """
        self._log("🔍 [A-Vision] Iniciando análisis visual de accesibilidad en lote...")

        if not rutas_fotos:
            self._log("   ⚠ No hay fotos para analizar.")
            return self._vision_default()

        # 1. Preparar las imágenes y redimensionarlas
        content_parts = [PROMPT_VISION_BATCH.format(marco_legal=MARCO_LEGAL)]
        rutas_validas = []

        for i, ruta in enumerate(rutas_fotos):
            if not os.path.exists(ruta):
                continue
            
            try:
                img = Image.open(ruta)
                img.thumbnail((800, 800))
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                content_parts.append(f"--- FOTO {len(rutas_validas)} ---")
                content_parts.append(img)
                rutas_validas.append(ruta)
                self._log(f"   📷 Cargada: {os.path.basename(ruta)}")
            except Exception as e:
                self._log(f"   ✗ Error procesando {ruta}: {e}")

        if not rutas_validas:
            self._log("   ⚠ Ninguna foto pudo ser leída correctamente.")
            return self._vision_default()

        # 2. Llamada única a Gemini con todas las fotos, desactivando filtros de seguridad
        self._log(f"   📡 Enviando {len(rutas_validas)} fotos simultáneamente a Gemini...")
        try:
            response = self.model.generate_content(
                content_parts,
                generation_config=genai.GenerationConfig(temperature=0.1, max_output_tokens=4096),
                safety_settings={
                    genai.types.HarmCategory.HARM_CATEGORY_HARASSMENT: genai.types.HarmBlockThreshold.BLOCK_NONE,
                    genai.types.HarmCategory.HARM_CATEGORY_HATE_SPEECH: genai.types.HarmBlockThreshold.BLOCK_NONE,
                    genai.types.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: genai.types.HarmBlockThreshold.BLOCK_NONE,
                    genai.types.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: genai.types.HarmBlockThreshold.BLOCK_NONE,
                }
            )
            
            texto = response.text.strip()
            
            # Extracción robusta de JSON
            start = texto.find('{')
            end = texto.rfind('}')
            if start != -1 and end != -1:
                texto = texto[start:end+1]
                
            resp_json = json.loads(texto)
            mapeo = resp_json.get("mapeo", {})
            analisis = resp_json.get("analisis", {})
            
            # Construir fotos_mapeadas
            fotos_mapeadas = {}
            # Fila 2 = Entrada
            try:
                idx_entrada = int(mapeo.get("entrada", -1))
                if idx_entrada != -1 and 0 <= idx_entrada < len(rutas_validas):
                    fotos_mapeadas[2] = rutas_validas[idx_entrada]
            except (ValueError, TypeError): pass
            
            # Fila 4 = Baños
            try:
                idx_banos = int(mapeo.get("banos", -1))
                if idx_banos != -1 and 0 <= idx_banos < len(rutas_validas):
                    fotos_mapeadas[4] = rutas_validas[idx_banos]
            except (ValueError, TypeError): pass
            
            # Fila 6 = Cafetería
            try:
                idx_cafeteria = int(mapeo.get("cafeteria", -1))
                if idx_cafeteria != -1 and 0 <= idx_cafeteria < len(rutas_validas):
                    fotos_mapeadas[6] = rutas_validas[idx_cafeteria]
            except (ValueError, TypeError): pass

            analisis["fotos_mapeadas"] = fotos_mapeadas
            self._log("   ✓ [A-Vision] Análisis visual en lote completado exitosamente.")
            return analisis

        except Exception as e:
            error_msg = str(e)
            self._log(f"   ✗ Error en ejecución en lote de A-Vision: {error_msg}")
            res_def = self._vision_default()
            res_def["acceso_principal"] = f"ERROR DE IA: {error_msg}. Por favor revise los logs o intente con menos fotos."
            res_def["fotos_mapeadas"] = {}
            return res_def

    def _vision_default(self):
        """Retorna análisis por defecto si falla la visión."""
        return {
            "acceso_principal": "No se pudo realizar el análisis visual automático. Requiere revisión manual.",
            "banos": "No se pudo realizar el análisis visual automático. Requiere revisión manual.",
            "cafeteria": "No se evidenció área de cafetería durante la visita.",
            "rampas": "No se pudo realizar el análisis visual automático. Requiere revisión manual.",
            "informacion_comunicacion": "No se pudo realizar el análisis visual automático. Requiere revisión manual.",
            "actitudinales_sociales": "No se pudo realizar el análisis visual automático. Requiere revisión manual.",
            "condiciones_biologico": "Requiere evaluación presencial.",
            "condiciones_psicosocial": "Requiere evaluación presencial.",
            "condiciones_biomecanico": "Requiere evaluación presencial.",
            "condiciones_fisico": "Requiere evaluación presencial.",
            "condiciones_quimico": "Requiere evaluación presencial.",
            "condiciones_seguridad": "Requiere evaluación presencial.",
        }

    # ── A-Legal ──────────────────────────────────────────────────────

    def ejecutar_legal(self, datos_cargo: dict, analisis_vision: dict) -> dict:
        """
        Ejecuta el Agente A-Legal para diagnóstico jurídico.

        Args:
            datos_cargo: Diccionario con info del cargo (nombre, tipo, objetivo, tareas).
            analisis_vision: Resultado del Agente A-Vision.

        Returns:
            dict con discapacidades compatibles, ajustes y conclusiones.
        """
        self._log("\n⚖️ [A-Legal] Iniciando diagnóstico jurídico de inclusión...")

        # Formatear tareas
        tareas_str = "\n".join(f"  - {t}" for t in datos_cargo.get("tareas", []))

        # Formatear análisis de visión
        vision_str = "\n".join(f"  • {k}: {v}" for k, v in analisis_vision.items())

        prompt = PROMPT_LEGAL.format(
            marco_legal=MARCO_LEGAL,
            cargo_nombre=datos_cargo.get("nombre", "Sin especificar"),
            cargo_tipo=datos_cargo.get("tipo", "Sin especificar"),
            cargo_objetivo=datos_cargo.get("objetivo", "Sin especificar"),
            tareas=tareas_str,
            analisis_vision=vision_str,
        )

        try:
            self._log("   📡 Consultando modelo de lenguaje para análisis jurídico...")
            response = self.model.generate_content(
                prompt,
                generation_config=genai.GenerationConfig(
                    temperature=0.2,
                    max_output_tokens=8192,
                ),
            )

            texto = response.text.strip()
            texto = re.sub(r'^```json\s*', '', texto)
            texto = re.sub(r'\s*```$', '', texto)

            resultado = json.loads(texto)
            self._log("   ✓ [A-Legal] Diagnóstico jurídico completado exitosamente.")
            return resultado

        except json.JSONDecodeError as e:
            self._log(f"   ✗ Error parseando respuesta JSON: {e}")
            self._log(f"   Respuesta raw: {response.text[:500]}...")
            return self._legal_default()
        except Exception as e:
            self._log(f"   ✗ Error en API de Gemini: {e}")
            return self._legal_default()

    def _legal_default(self):
        """Retorna análisis legal por defecto si falla."""
        return {
            "discapacidades_compatibles": ["Auditiva", "Intelectual Leve"],
            "ajustes_razonables": [
                {
                    "tipo": "Auditiva",
                    "fisicas": "Requiere análisis manual por profesional.",
                    "comunicacion": "Requiere análisis manual por profesional.",
                    "actitudinales": "Requiere análisis manual por profesional.",
                },
            ],
            "competencias": ["Medio – Requiere evaluación manual."] * 35,
            "conclusiones": "El análisis automático no pudo completarse. Se requiere revisión por un profesional de inclusión laboral.",
        }

    # ── Orquestación Completa ────────────────────────────────────────

    def ejecutar_analisis_completo(
        self,
        rutas_fotos: list,
        datos_cargo: dict,
        progress_callback=None,
    ) -> dict:
        """
        Ejecuta el pipeline completo: A-Vision → A-Legal → Resultado.

        Args:
            rutas_fotos: Lista de rutas a las fotos.
            datos_cargo: Dict con nombre, tipo, objetivo, tareas del cargo.
            progress_callback: Función callback(fase, progreso) opcional.

        Returns:
            dict con toda la información generada por los agentes.
        """
        cb = progress_callback or (lambda f, p: None)

        self._log("═" * 60)
        self._log("  SISTEMA MULTIAGENTE DE INCLUSIÓN LABORAL")
        self._log("  Comfenalco Santander — Talento Sin Barreras")
        self._log("═" * 60)

        # Fase 1: Análisis Visual
        cb("Fase 1: Análisis Visual (A-Vision)...", 0.2)
        analisis_vision = self.ejecutar_vision(rutas_fotos)

        # Fase 2: Diagnóstico Jurídico
        cb("Fase 2: Diagnóstico Jurídico (A-Legal)...", 0.5)
        analisis_vision_clean = {k: v for k, v in analisis_vision.items() if k != "fotos_mapeadas"}
        analisis_legal = self.ejecutar_legal(datos_cargo, analisis_vision_clean)

        # Fase 3: Compilar resultados
        cb("Fase 3: Compilando resultados...", 0.8)

        # Construir la estructura de accesibilidad (9 filas para Tabla 5)
        accesibilidad = [
            analisis_vision.get("acceso_principal", ""),
            "",  # Fila de foto
            analisis_vision.get("banos", ""),
            "",  # Fila de foto
            analisis_vision.get("cafeteria", ""),
            "No se cuenta con registro fotográfico de cafetería.",
            analisis_vision.get("rampas", ""),
            analisis_vision.get("informacion_comunicacion", ""),
            analisis_vision.get("actitudinales_sociales", ""),
        ]

        # Condiciones ambientales
        condiciones_amb = [
            analisis_vision.get("condiciones_biologico", ""),
            analisis_vision.get("condiciones_psicosocial", ""),
            analisis_vision.get("condiciones_biomecanico", ""),
            analisis_vision.get("condiciones_fisico", ""),
            analisis_vision.get("condiciones_quimico", ""),
        ]

        condiciones_seg_texto = analisis_vision.get("condiciones_seguridad", "")
        condiciones_seguridad = [
            condiciones_seg_texto,
            "No aplica",
            condiciones_seg_texto,
            "Matriz de identificación de peligros y valoración de riesgos (GTC 45)",
        ]

        resultado = {
            "accesibilidad": accesibilidad,
            "condiciones_amb": condiciones_amb,
            "condiciones_seguridad": condiciones_seguridad,
            "competencias": analisis_legal.get("competencias", ["Medio"] * 35),
            "discapacidades_sugeridas": analisis_legal.get("discapacidades_compatibles", []),
            "ajustes_razonables": analisis_legal.get("ajustes_razonables", []),
            "conclusiones": analisis_legal.get("conclusiones", ""),
            "fotos_mapeadas": analisis_vision.get("fotos_mapeadas", {})
        }

        cb("Análisis completo finalizado.", 1.0)
        self._log("\n✓ Sistema multiagente: Análisis completo finalizado.")
        self._log("═" * 60)

        return resultado


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  UTILIDADES                                                                ║
# ╚══════════════════════════════════════════════════════════════════════════════╝


def extraer_texto_docx(ruta_docx: str) -> str:
    """
    Extrae todo el texto de un archivo .docx (párrafos, tablas, encabezados y pies).
    """
    from docx import Document

    doc = Document(ruta_docx)
    textos = []

    # Extraer encabezados y pies de página
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                if p.text.strip():
                    textos.append(p.text.strip())
            for t in section.header.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            textos.append(cell.text.strip())
        if section.footer:
            for p in section.footer.paragraphs:
                if p.text.strip():
                    textos.append(p.text.strip())
            for t in section.footer.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            textos.append(cell.text.strip())

    # Cuerpo principal
    for p in doc.paragraphs:
        if p.text.strip():
            textos.append(p.text.strip())

    for tabla in doc.tables:
        for row in tabla.rows:
            for cell in row.cells:
                if cell.text.strip():
                    textos.append(cell.text.strip())

    return "\n".join(textos)


def extraer_texto_pdf(ruta_pdf: str) -> str:
    try:
        import pypdf
        textos = []
        with open(ruta_pdf, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    textos.append(t.strip())
        return "\n".join(textos)
    except Exception as e:
        return f"Error extrayendo PDF: {e}"


def extraer_texto_excel(ruta_excel: str) -> str:
    try:
        import pandas as pd
        textos = []
        xls = pd.ExcelFile(ruta_excel)
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            textos.append(f"--- Hoja: {sheet_name} ---")
            textos.append(df.to_string(index=False))
        return "\n".join(textos)
    except Exception as e:
        return f"Error extrayendo Excel: {e}"


def extraer_texto_plano(ruta_txt: str) -> str:
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            with open(ruta_txt, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    return ""


def extraer_texto_de_archivo(ruta: str) -> str:
    ext = Path(ruta).suffix.lower()
    if ext == ".docx":
        return extraer_texto_docx(ruta)
    elif ext == ".pdf":
        return extraer_texto_pdf(ruta)
    elif ext in [".xlsx", ".xls"]:
        return extraer_texto_excel(ruta)
    elif ext in [".txt", ".csv", ".json"]:
        return extraer_texto_plano(ruta)
    elif ext == ".doc":
        try:
            return extraer_texto_docx(ruta)
        except Exception:
            pass
        try:
            with open(ruta, "rb") as f:
                content = f.read()
            import re
            strings = re.findall(rb'[a-zA-Z0-9\s\.,;:!\?\-\(\)\[\]\{\}\/\\_@#\$%&\*\+\=\<\>\'\"]{4,}', content)
            text = ""
            for s in strings:
                try:
                    text += s.decode('utf-8') + "\n"
                except Exception:
                    try:
                        text += s.decode('latin-1') + "\n"
                    except Exception:
                        pass
            return text
        except Exception as e:
            return f"No se pudo extraer texto del archivo .doc: {e}"
    else:
        try:
            return extraer_texto_plano(ruta)
        except Exception as e:
            return f"Formato no soportado: {e}"


def extraer_datos_manual(ruta_manual: str, api_key: str = None) -> dict:
    """
    Extrae datos del Manual de Funciones usando Gemini.
    Soporta docx, pdf, xls, txt, doc.
    """
    texto = extraer_texto_de_archivo(ruta_manual)

    if api_key and GEMINI_DISPONIBLE:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(MODELO_GEMINI)

        prompt = f"""
Analiza el siguiente texto de un Manual de Funciones/Descripción de Cargo y extrae la información en formato JSON. Responde ÚNICAMENTE con el JSON (sin markdown ni ```json).
Si un dato no existe en el manual, déjalo como una cadena vacía "".

TEXTO DEL MANUAL:
{texto[:12000]}

JSON requerido:
{{
    "empresa_nombre": "Nombre de la empresa",
    "empresa_actividad": "Actividad económica",
    "empresa_nit": "NIT",
    "empresa_sede": "Sede o Ciudad",
    "empresa_direccion": "Dirección",
    "empresa_telefono": "Teléfono",
    "empresa_fecha": "Fecha de actualización o elaboración",
    "cargo_nombre": "Nombre del cargo",
    "cargo_tipo": "Tipo (Operativo/Administrativo/etc.)",
    "cargo_dependencia": "Dependencia o área",
    "cargo_reporta": "Cargo al que reporta",
    "cargo_objetivo": "Objetivo principal del cargo",
    "cargo_requisito_educativo": "Nivel educativo requerido",
    "cargo_requisito_certificaciones": "Certificaciones requeridas",
    "cargo_requisito_conocimientos": "Conocimientos específicos",
    "cargo_requisito_experiencia": "Experiencia requerida",
    "cargo_requisito_entrenamiento": "Entrenamiento requerido",
    "cargo_requisito_disponibilidad": "Disponibilidad (turnos, viajes, etc.)",
    "cargo_requisito_examenes": "Exámenes médicos",
    "cargo_tareas": ["Tarea 1", "Tarea 2", "Tarea 3", "Tarea 4", "Tarea 5"],
    "condiciones_jornada": "Jornada laboral",
    "condiciones_turnos": "Turnos",
    "condiciones_rotativos": "Turnos rotativos",
    "condiciones_rotacion": "Rotación de puestos",
    "condiciones_horas_extras": "Horas extras",
    "recursos_equipos": "Equipos de oficina/cómputo",
    "recursos_mobiliario": "Mobiliario",
    "recursos_maquinas": "Máquinas/vehículos",
    "recursos_herramientas": "Herramientas manuales",
    "recursos_materiales": "Materiales e insumos",
    "recursos_accesorios": "Accesorios de dotación",
    "recursos_epp": "Elementos de Protección Personal (EPP)"
}}
"""
        try:
            response = model.generate_content(
                prompt,
                generation_config=genai.GenerationConfig(
                    temperature=0.1,
                    max_output_tokens=8192,
                ),
            )
            texto_resp = response.text.strip()
            texto_resp = re.sub(r'^```json\s*', '', texto_resp)
            texto_resp = re.sub(r'\s*```$', '', texto_resp)
            return json.loads(texto_resp)
        except Exception:
            pass

    return {"texto_crudo": texto}


def verificar_api_key(api_key: str) -> tuple:
    """Verifica que la API key de Gemini sea válida y retorna (es_valida, mensaje_error)."""
    if not GEMINI_DISPONIBLE:
        return False, "La librería 'google-generativeai' no está disponible en el servidor."
    if not api_key:
        return False, "La clave API está vacía."
    try:
        # Limpiar posibles espacios en blanco alrededor de la clave
        api_key = api_key.strip()
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(MODELO_GEMINI)
        response = model.generate_content(
            "Responde solo 'OK'.",
            generation_config=genai.GenerationConfig(max_output_tokens=10),
        )
        if response.text:
            return True, "API Key válida"
        return False, "El modelo no retornó ninguna respuesta."
    except Exception as e:
        error_msg = str(e)
        # Simplificar mensajes de error comunes de Google
        if "API_KEY_INVALID" in error_msg:
            return False, "La API Key ingresada no es válida para Google Cloud."
        elif "quota" in error_msg.lower() or "limit" in error_msg.lower():
            return False, "Se ha excedido la cuota o límite de tu API Key."
        return False, f"Error al conectar con Google Gemini: {error_msg}"

